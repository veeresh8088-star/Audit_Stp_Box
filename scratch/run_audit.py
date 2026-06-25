import os
import sys
import zipfile
import io
import json
import easyocr
import PIL.Image
import numpy as np

# Add workspace directory to python path
sys.path.append(os.getcwd())

from src.db.database import init_db, SessionLocal, DocumentChunk, force_master
from src.core.controls_data import USE_CASES
from src.ai.audit_graph import audit_graph
from src.core.retrieval import save_document_chunks, _ingested_chunks_cache

def chunk_paragraphs(paragraphs_data, target=1000, overlap=200):
    chunks = []
    if not paragraphs_data:
        return []
    current_chunk_paras = []
    current_len = 0
    idx = 0
    while idx < len(paragraphs_data):
        p_text, section = paragraphs_data[idx]
        current_chunk_paras.append((p_text, section, idx))
        current_len += len(p_text) + 2
        
        if current_len >= target or idx == len(paragraphs_data) - 1:
            chunk_section = ""
            for _, sec, _ in current_chunk_paras:
                if sec:
                    chunk_section = sec
            chunk_content = "\n\n".join([txt for txt, _, _ in current_chunk_paras])
            chunks.append((chunk_content, chunk_section, current_chunk_paras[0][2], current_chunk_paras[-1][2]))
            
            overlap_len = 0
            overlap_paras = []
            for txt, sec, p_idx in reversed(current_chunk_paras):
                if overlap_len + len(txt) + 2 <= overlap or not overlap_paras:
                    overlap_paras.insert(0, (txt, sec, p_idx))
                    overlap_len += len(txt) + 2
                else:
                    break
            if len(overlap_paras) == len(current_chunk_paras):
                if len(overlap_paras) > 1:
                    overlap_paras = overlap_paras[1:]
                else:
                    overlap_paras = []
            current_chunk_paras = list(overlap_paras)
            current_len = sum(len(txt) + 2 for txt, _, _ in current_chunk_paras)
        idx += 1
    return chunks

def extract_text_and_chunks(doc_path):
    print(f"Loading document: {doc_path}")
    doc = Document(doc_path) if 'Document' in globals() else None
    
    # Import inside if needed
    from docx import Document
    doc = Document(doc_path)
    
    paragraphs_data = []
    
    # 1. Paragraphs with heading detection
    current_section = ""
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if p_text:
            if p.style and p.style.name and p.style.name.startswith("Heading"):
                current_section = p_text
            paragraphs_data.append((p_text, current_section))
            
    # 2. Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells_text:
                deduped_cells = []
                for cell_txt in row_cells_text:
                    if not deduped_cells or deduped_cells[-1] != cell_txt:
                        deduped_cells.append(cell_txt)
                if deduped_cells:
                    row_text = " | ".join(deduped_cells)
                    paragraphs_data.append((row_text, "[Table Data]"))
                    
    # 3. Extract and OCR images from ZIP
    try:
        with zipfile.ZipFile(doc_path) as zf:
            media_files = [n for n in zf.namelist() if n.startswith("word/media/") and n.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".tiff", ".bmp"))]
            if media_files:
                print("Performing OCR on embedded images...")
                reader = easyocr.Reader(['en'], gpu=False)
                for name in sorted(media_files):
                    try:
                        img_data = zf.read(name)
                        img = PIL.Image.open(io.BytesIO(img_data))
                        img_np = np.array(img)
                        res = reader.readtext(img_np, detail=0)
                        if res:
                            ocr_text = " ".join(res)
                            base_img_name = os.path.basename(name)
                            paragraphs_data.append((f"[Embedded Image OCR ({base_img_name})]: {ocr_text}", "[Embedded Image Content]"))
                            print(f"OCR succeeded for {base_img_name}: {ocr_text[:80]}...")
                    except Exception as img_err:
                        print("Error on image OCR:", img_err)
    except Exception as zf_err:
        print("Error reading zip:", zf_err)
        
    docx_fname = os.path.basename(doc_path)
    docx_ext = "docx"
    
    chunks = chunk_paragraphs(paragraphs_data, target=1000, overlap=200)
    docx_chunks = []
    for chunk_content, chunk_section, _, _ in chunks:
        docx_chunks.append((chunk_content, {
            "source_file": docx_fname,
            "source_type": docx_ext,
            "section_heading": chunk_section,
            "chunk_id": ""
        }))
        
    _ingested_chunks_cache[docx_fname] = docx_chunks
    raw_text = "\n\n".join([txt for txt, _ in paragraphs_data])
    return raw_text, docx_fname

def run_audits():
    doc_path = "10 -Multi-factor authentication operator.docx"
    raw_text, docx_fname = extract_text_and_chunks(doc_path)
    
    print("Saving document chunks to database...")
    save_document_chunks(docx_fname, raw_text)
    
    # Verify chunks in database
    db = SessionLocal()
    cnt = db.query(DocumentChunk).filter(DocumentChunk.filename == docx_fname).count()
    db.close()
    print(f"Verified: {cnt} chunks saved in database.")
    
    # Locate controls
    target_ids = ["8.5", "8.2", "8.15"]
    target_controls = []
    for c_id in target_ids:
        found = None
        for c in USE_CASES:
            # Match start of use case e.g. "8.5 "
            if c['use_case'].split(' ')[0] == c_id:
                found = c
                break
        if found:
            target_controls.append(found)
        else:
            print(f"Warning: Control {c_id} not found in USE_CASES")
            
    print(f"Found {len(target_controls)} matching controls for audit.")
    
    results = {}
    
    for idx, control in enumerate(target_controls):
        print(f"\n--- Auditing control {control['use_case']} using qwen2.5:7b ---")
        state = {
            "control_id": control["use_case"],
            "control_label": control["label"],
            "expected_evidence": control["expected"],
            "prompt_hint": control.get("prompt_hint", ""),
            "severity": control["severity"],
            "standard": control.get("standard", "ISO 27001"),
            "recommendation": control.get("recommendation", ""),
            
            "document_text": raw_text,
            "file_names_list": [docx_fname],
            "ollama_model": "qwen2.5:7b",
            "summary_text": f"AWS IAM login screenshot demonstration demonstrating MFA configuration settings for operator Rakesh_Kumar_Sharma.",
            
            "retrieved_context": "",
            "draft_finding": None,
            "validation_error": None,
            "retry_count": 0,
            "final_finding": None,
            
            "bg_key": "cmd-audit",
            "control_idx": idx,
            "total_controls": len(target_controls)
        }
        
        try:
            output_state = audit_graph.invoke(state)
            final_finding = output_state.get("final_finding")
            results[control["use_case"]] = final_finding
            print(f"Result Status: {final_finding.get('status') if final_finding else 'None'}")
        except Exception as e:
            print(f"Error auditing control {control['use_case']}: {e}")
            results[control["use_case"]] = {"status": "ERROR", "justification": str(e)}
            
    # Save results to JSON file
    with open("scratch/audit_results.json", "w", encoding="utf-8") as f:
        json.dump(results, f, indent=4)
    print("\nSaved results to scratch/audit_results.json")
    
if __name__ == "__main__":
    run_audits()
